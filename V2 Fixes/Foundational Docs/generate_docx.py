#!/usr/bin/env python3
"""Generate Word documents from all v2 deliverables."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE = "/Users/p00l/Desktop/Upgrading Foundational Docs/Upgraded Prompts"
OUTPUT_DIR = os.path.join(BASE, "Word Documents")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Clear old Word docs first
for f in os.listdir(OUTPUT_DIR):
    if f.endswith('.docx'):
        os.remove(os.path.join(OUTPUT_DIR, f))

FILES = {
    # Research pipeline prompts
    "01_Competitor_Research_v2": os.path.join(BASE, "clean_prompts/01_competitor_research_v2.md"),
    "03_Deep_Research_Meta_Prompt_v2": os.path.join(BASE, "clean_prompts/03_deep_research_meta_prompt_v2.md"),
    "04_Deep_Research_Execution_v2": os.path.join(BASE, "clean_prompts/04_deep_research_execution_v2.md"),
    "06_Avatar_Brief_v2": os.path.join(BASE, "clean_prompts/06_avatar_brief_v2.md"),
    # Pipeline overview
    "00_Pipeline_Overview": os.path.join(BASE, "pipeline/00_pipeline_overview.md"),
    # Downstream prompts (not part of research pipeline)
    "07_Offer_Brief_v2_DOWNSTREAM": os.path.join(BASE, "downstream/07_offer_brief_v2.md"),
    "08_Belief_Architecture_v2_DOWNSTREAM": os.path.join(BASE, "downstream/08_belief_architecture_v2.md"),
    # Testing report
    "Testing_Report_v1_vs_v2": os.path.join(BASE, "reference/docs/testing-report-v1-vs-v2.md"),
}


def setup_styles(doc):
    """Configure document styles for professional appearance."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        heading_font.bold = True
        if level == 1:
            heading_font.size = Pt(22)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading_font.size = Pt(16)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        elif level == 3:
            heading_font.size = Pt(13)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        else:
            heading_font.size = Pt(11)
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)


def parse_table(lines, start_idx):
    rows = []
    idx = start_idx
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[\s\-\|:]+\|$', line):
            idx += 1
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        idx += 1
    return rows, idx


def add_table_to_doc(doc, rows):
    if not rows or len(rows) < 1:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = clean_markdown(cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.bold = True


def clean_markdown(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


def add_formatted_paragraph(doc, text, is_bold=False):
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*') and not sub.startswith('**'):
                    run = p.add_run(sub[1:-1])
                    run.italic = True
                else:
                    clean = clean_markdown(sub)
                    if clean:
                        run = p.add_run(clean)
    if is_bold:
        for run in p.runs:
            run.bold = True
    return p


def md_to_docx(md_path, output_name):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()
    setup_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code_block:
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    p.paragraph_format.left_indent = Inches(0.3)
                    code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run('_' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
            i += 1
            continue

        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)', stripped)
            if match:
                level = len(match.group(1))
                header_text = clean_markdown(match.group(2))
                doc.add_heading(header_text, level=min(level, 4))
                i += 1
                continue

        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            rows, end_idx = parse_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows)
                doc.add_paragraph()
            i = end_idx
            continue

        if re.match(r'^[\s]*[-*]\s', stripped) or re.match(r'^[\s]*\d+\.\s', stripped):
            leading_spaces = len(line) - len(line.lstrip())
            indent_level = leading_spaces // 2
            bullet_text = re.sub(r'^[\s]*[-*]\s+', '', stripped)
            bullet_text = re.sub(r'^[\s]*\d+\.\s+', '', stripped)
            p = add_formatted_paragraph(doc, bullet_text)
            p.style = 'List Bullet' if re.match(r'^[\s]*[-*]\s', stripped) else 'List Number'
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            i += 1
            continue

        if stripped.startswith('>'):
            quote_text = stripped.lstrip('> ').strip()
            p = add_formatted_paragraph(doc, quote_text)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1

    output_path = os.path.join(OUTPUT_DIR, f"{output_name}.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path


if __name__ == "__main__":
    for name, path in FILES.items():
        if os.path.exists(path):
            md_to_docx(path, name)
        else:
            print(f"SKIP (not found): {path}")

    print(f"\nAll documents saved to: {OUTPUT_DIR}")
